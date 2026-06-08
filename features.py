"""
增强功能模块: 录取概率 + 志愿方案生成 + 院校对比 + PDF报告
"""

import os, datetime
from database import get_db

# ============================================================
# 1. 录取概率估算
# ============================================================
def estimate_probability(composite, batch):
    """将综合评分映射为录取概率百分比"""
    if composite >= 90:
        return 95, "极高"
    elif composite >= 80:
        return 88, "很高"
    elif composite >= 70:
        return 75, "较高"
    elif composite >= 60:
        return 60, "中等"
    elif composite >= 50:
        return 45, "偏低"
    elif composite >= 40:
        return 30, "较低"
    elif composite >= 30:
        return 18, "很低"
    else:
        return 8, "极低"

# ============================================================
# 2. 志愿方案生成器
# ============================================================
def generate_plan(result, province, category, score):
    """从推荐结果中自动组合最优志愿方案"""
    plan = {
        "province": province,
        "category": category,
        "score": score,
        "冲": [], "稳": [], "保": [],
        "total_schools": 0,
    }
    
    # 冲: 取综合分最低的3所(最难但最值得冲刺的)
    chong_candidates = sorted(result.get("冲", []), key=lambda x: x["composite"])
    for s in chong_candidates[:3]:
        prob, label = estimate_probability(s["composite"], result["my_info"]["batch"])
        plan["冲"].append({
            "name": s["name"], "level": s["level"], "city": s["city"],
            "score": s["uni_avg_score"], "composite": s["composite"],
            "probability": prob, "prob_label": label,
            "majors": s.get("majors_bao", [])[:3] if s.get("majors_bao") else [],
        })
    
    # 稳: 取综合分中等的5所
    wen_candidates = sorted(result.get("稳", []), key=lambda x: x["composite"])
    mid = len(wen_candidates) // 2
    start = max(0, mid - 3)
    for s in wen_candidates[start:start+5]:
        prob, label = estimate_probability(s["composite"], result["my_info"]["batch"])
        plan["稳"].append({
            "name": s["name"], "level": s["level"], "city": s["city"],
            "score": s["uni_avg_score"], "composite": s["composite"],
            "probability": prob, "prob_label": label,
            "majors": s.get("majors_bao", [])[:3] if s.get("majors_bao") else [],
        })
    
    # 保: 取综合分最高的3所
    bao_candidates = sorted(result.get("保", []), key=lambda x: x["composite"], reverse=True)
    for s in bao_candidates[:3]:
        prob, label = estimate_probability(s["composite"], result["my_info"]["batch"])
        plan["保"].append({
            "name": s["name"], "level": s["level"], "city": s["city"],
            "score": s["uni_avg_score"], "composite": s["composite"],
            "probability": prob, "prob_label": label,
            "majors": s.get("majors_bao", [])[:3] if s.get("majors_bao") else [],
        })
    
    plan["total_schools"] = len(plan["冲"]) + len(plan["稳"]) + len(plan["保"])
    return plan

# ============================================================
# 3. 院校对比数据
# ============================================================
def get_school_detail(school_name):
    """获取院校详细信息用于对比"""
    conn = get_db()
    uni = conn.execute("SELECT * FROM universities WHERE name=?", (school_name,)).fetchone()
    if not uni:
        conn.close()
        return None
    
    # 获取录取数据
    admissions = conn.execute("""
        SELECT province, year, category, batch, min_score, min_rank, avg_score
        FROM admission_scores WHERE university_id=?
        ORDER BY year DESC, province
    """, (uni['id'],)).fetchall()
    
    # 获取优势专业
    majors = conn.execute("""
        SELECT m.name, m.category, m.employment_score, m.avg_salary
        FROM uni_majors um JOIN majors m ON um.major_id=m.id
        WHERE um.university_id=? AND um.is_advantage=1
        ORDER BY m.employment_score DESC LIMIT 5
    """, (uni['id'],)).fetchall()
    
    conn.close()
    
    # 按省份聚合录取分
    prov_scores = {}
    for a in admissions:
        if a['province'] not in prov_scores:
            prov_scores[a['province']] = []
        prov_scores[a['province']].append({
            "year": a['year'], "category": a['category'],
            "score": a['min_score'], "rank": a['min_rank']
        })
    
    return {
        "name": uni['name'], "level": uni['level'], "type": uni['type'],
        "city": uni['city'], "is_public": bool(uni['is_public']),
        "majors": [dict(m) for m in majors],
        "admissions": prov_scores,
    }

# ============================================================
# 4. PDF报告生成
# ============================================================
def generate_pdf_report(plan, user_info, output_path):
    """生成志愿填报方案PDF报告"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    
    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(60)
    run = p.add_run('高考志愿填报方案')
    run.bold = True; run.font.size = Pt(24); run.font.color.rgb = RGBColor(0, 102, 204)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{plan['province']} · {plan['category']} · {plan['score']}分")
    run2.font.size = Pt(14); run2.font.color.rgb = RGBColor(100,100,100)
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_after = Pt(40)
    date_str = datetime.date.today().strftime("%Y年%m月%d日")
    run3 = p3.add_run(f'生成日期: {date_str}\n本报告由AI辅助生成，仅供参考')
    run3.font.size = Pt(9); run3.font.color.rgb = RGBColor(150,150,150)
    
    # 免责
    doc.add_heading('⚠️ 重要提示', level=2)
    doc.add_paragraph(
        '本报告基于历史公开数据提供参考性志愿填报建议，不构成任何报考决策依据。'
        '最终志愿填报请以各省教育考试院官方发布信息为准。'
        '使用本报告即表示您已理解并接受相关风险。'
    )
    
    # 考生信息
    doc.add_heading('一、考生信息', level=1)
    info = user_info
    info_data = [
        ['省份', info.get('province','')], ['科类', info.get('category','')],
        ['分数', f"{info.get('score','')}分"], ['位次', f"{info.get('rank',0):,}"],
        ['批次', info.get('batch','')], ['批次线', f"{info.get('batch_line','')}分"],
        ['线差', f"+{info.get('score',0) - info.get('batch_line',0)}分"],
        ['报告日期', date_str],
    ]
    table = doc.add_table(rows=len(info_data)+1, cols=2, style='Light Grid Accent 1')
    for i, h in enumerate(['项目','内容']):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs: run.bold = True
    for i, row in enumerate(info_data):
        for j, txt in enumerate(row):
            table.rows[i+1].cells[j].text = str(txt)
    
    # 推荐方案
    for tag, label, desc in [("冲","二、冲刺院校（可争取）","录取概率较低，但值得冲刺的院校"),
                              ("稳","三、稳妥院校（较匹配）","与你的位次匹配度较高"),
                              ("保","四、保底院校（稳录取）","录取概率很高，确保有学上")]:
        doc.add_heading(label, level=1)
        doc.add_paragraph(desc)
        
        if not plan.get(tag):
            doc.add_paragraph('（无推荐）')
            continue
        
        schools_data = []
        for s in plan[tag]:
            schools_data.append([
                s['name'], s['level'], s['city'],
                f"{s['score']}分", f"{s['composite']:.0f}分",
                f"{s.get('probability','?')}%", 
                ', '.join(s.get('majors',[])) if s.get('majors') else '-'
            ])
        
        t = doc.add_table(rows=len(schools_data)+1, cols=7, style='Light Grid Accent 1')
        headers = ['院校','层次','城市','录取分','综合分','录取概率','推荐专业']
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
            for run in t.rows[0].cells[i].paragraphs[0].runs: run.bold = True; run.font.size = Pt(8)
        for i, row in enumerate(schools_data):
            for j, txt in enumerate(row):
                t.rows[i+1].cells[j].text = txt
                for run in t.rows[i+1].cells[j].paragraphs[0].runs: run.font.size = Pt(8)
    
    # 建议
    doc.add_heading('五、使用建议', level=1)
    tips = [
        '冲(2-3所): 选最喜欢的2-3所冲刺校，不要全填冲刺',
        '稳(3-5所): 选3-5所匹配校，这是最可能被录取的区间',
        '保(2-3所): 选2-3所保底校，确保万无一失',
        '最终志愿需结合专业兴趣、城市偏好、招生计划等因素综合决策',
        '请务必核对各省教育考试院官方发布的招生计划和录取数据',
    ]
    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    # 页脚
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(180,180,180); run.font.size = Pt(8)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f'AI辅助生成 · {date_str} · 仅供参考，不构成报考建议')
    run2.font.size = Pt(8); run2.font.color.rgb = RGBColor(150,150,150)
    
    doc.save(output_path)
    return output_path
